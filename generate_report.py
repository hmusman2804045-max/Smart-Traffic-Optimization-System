import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_heading_style(paragraph, font_size=16, bold=True):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.bold = bold

def create_report():
    doc = docx.Document()

    # --- TITLE PAGE ---
    # Section Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deep Learning Section V-3")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("\n" * 2)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Based Intelligent Traffic Optimization System")
    run.font.size = Pt(24)
    run.bold = True

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Final Project")
    run.font.size = Pt(18)
    run.italic = True

    doc.add_paragraph("\n" * 5)

    # Submission Details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Submitted by:")
    run.bold = True
    doc.add_paragraph("-------------------------") # Placeholder for name

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Submitted to:")
    run.bold = True
    doc.add_paragraph("Dr. Jameel Ahmad")

    doc.add_page_break()

    # --- INTRODUCTION ---
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "Modern urban environments face critical challenges due to traffic congestion, which leads to increased commute times, "
        "higher fuel consumption, and significant environmental impact. Traditional traffic management systems rely on fixed-time "
        "signals that do not adapt to real-time traffic flux. This project introduces an AI-driven solution that leverages "
        "multiple deep learning paradigms to optimize traffic flow dynamically."
    )

    doc.add_heading('2. Problem Statement', level=2)
    doc.add_paragraph(
        "Static traffic control systems are inefficient because they cannot respond to sudden surges in vehicle density or "
        "unexpected road incidents. There is a need for an integrated framework that can perceive current conditions, "
        "predict future trends, detect anomalies, and make intelligent decisions in real-time."
    )

    # --- SYSTEM ARCHITECTURE ---
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph(
        "The proposed system follows a multi-modal architecture where diverse data sources (visual, temporal, and textual) "
        "are processed by specialized deep learning models. These outputs are then fed into a Reinforcement Learning agent "
        "that executes the optimal traffic signal control policy."
    )

    # --- MODEL BREAKDOWN ---
    doc.add_heading('4. Core Modules and Model Details', level=1)

    # Perception
    doc.add_heading('4.1 Traffic Perception (YOLOv8)', level=2)
    doc.add_paragraph(
        "Purpose: The perception module acts as the 'eyes' of the system. It uses the YOLOv8 (You Only Look Once) architecture "
        "for high-speed, real-time object detection. The model identifies and counts vehicles (cars, trucks, buses, motorcycles) "
        "and pedestrians from live camera feeds.\n"
        "Technical Detail: Utilizing a pre-trained YOLOv8 nano model ensures low latency, making it suitable for edge deployment "
        "at intersections.\n"
        "Improvements: Future iterations can be fine-tuned on region-specific datasets like BDD100K to improve detection "
        "reliability under challenging conditions such as heavy rain, fog, or night-time environments."
    )

    # Prediction
    doc.add_heading('4.2 Traffic Flow Prediction (LSTM)', level=2)
    doc.add_paragraph(
        "Purpose: To move from reactive to proactive control, the system predicts upcoming traffic density. "
        "The Long Short-Term Memory (LSTM) network analyzes historical vehicle counts to forecast trends.\n"
        "Technical Detail: LSTMs are used because they can capture long-term temporal dependencies in time-series data, "
        "recognizing morning and evening rush hour cycles.\n"
        "Improvements: Accuracy can be enhanced by incorporating external variables such as public holiday calendars, "
        "major city events, and weather forecasts into the feature vector."
    )

    # Anomaly Detection
    doc.add_heading('4.3 Anomaly Detection (VAE)', level=2)
    doc.add_paragraph(
        "Purpose: Identifying erratic traffic behavior (e.g., accidents or vehicle breakdowns) is crucial for emergency response. "
        "A Variational Autoencoder (VAE) is trained to learn the 'latent representation' of normal traffic patterns.\n"
        "Technical Detail: When the reconstruction error exceeds a predefined threshold, the system flags a potential anomaly.\n"
        "Improvements: Implementing a dynamic threshold that adjusts based on the time of day can reduce false positives during "
        "naturally low-traffic periods (like 3 AM)."
    )

    # NLP
    doc.add_heading('4.4 Social Media Sentiment Analysis (BERT)', level=2)
    doc.add_paragraph(
        "Purpose: Human-reported data often precedes sensor-detected anomalies. This module mines social media feeds (e.g., Twitter/X) "
        "for traffic reports using a BERT-based NLP model.\n"
        "Technical Detail: The model classifies text into categories: Normal, Accident, Roadwork, or Weather-related incidents.\n"
        "Improvements: Future versions could include geo-tagging to pinpoint incidents more accurately and sentiment-based weighting "
        "where 'panic' or 'emergency' keywords trigger higher priority signals."
    )

    # --- REINFORCEMENT LEARNING ---
    doc.add_heading('5. Intelligent Decision Making (PPO RL)', level=1)
    doc.add_paragraph(
        "The brain of the system is a Reinforcement Learning agent using the Proximal Policy Optimization (PPO) algorithm. "
        "The agent receives a state vector containing:\n"
        "• Current vehicle density (from Perception)\n"
        "• Predicted traffic volume (from LSTM)\n"
        "• Anomaly flags (from VAE)\n"
        "• Sentiment and incident reports (from BERT)\n"
        "The agent is rewarded for minimizing cumulative traffic density while prioritizing clearing roads where anomalies "
        "or accidents are reported."
    )

    # --- CONCLUSION ---
    doc.add_heading('6. Conclusion and Future Scope', level=1)
    doc.add_paragraph(
        "This integrated system demonstrates the power of multi-modal Deep Learning in solving complex urban problems. "
        "By combining vision, time-series, and text data, the system provides a robust and adaptive traffic management framework.\n"
        "Future Scope: The next phase involves multi-intersection coordination, where RL agents communicate with each other "
        "to prevent congestion from simply moving to the next block, creating a 'Smart City' traffic grid."
    )

    # Save the document
    file_name = "Final_Project_Report.docx"
    doc.save(file_name)
    print(f"Report generated successfully: {file_name}")

if __name__ == "__main__":
    create_report()
